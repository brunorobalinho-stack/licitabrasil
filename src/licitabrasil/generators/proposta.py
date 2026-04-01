"""PropostaGenerator — gera documentos DOCX para licitacoes.

Usa docxtpl (python-docx + Jinja2) para gerar propostas de preco
e declaracoes a partir de templates .docx.

Templates base:
- Proposta de precos (planilha de custos e formacao de precos)
- Declaracao ME/EPP, fato impeditivo, menor, elaboracao independente
- Planilha de custos IN SEGES/ME 65/2021 (servicos continuados)
"""

from __future__ import annotations

import logging
from datetime import datetime
from decimal import Decimal
from pathlib import Path

from docxtpl import DocxTemplate
from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).parent / "templates"


# ---------------------------------------------------------------------------
# Models de entrada — Empresa e Licitacao
# ---------------------------------------------------------------------------

class ItemProposta(BaseModel):
    """Item individual da proposta de precos."""

    numero: int
    descricao: str
    unidade: str = "UN"
    quantidade: Decimal = Decimal("1")
    valor_unitario: Decimal = Decimal("0")

    @property
    def valor_total(self) -> Decimal:
        return self.quantidade * self.valor_unitario


class DadosEmpresa(BaseModel):
    """Dados da empresa para geracao de documentos."""

    razao_social: str
    cnpj: str
    endereco: str = ""
    telefone: str = ""
    email: str = ""
    representante_legal: str = ""
    cargo_representante: str = ""
    dados_bancarios: dict = Field(default_factory=dict)


class DadosLicitacao(BaseModel):
    """Dados da licitacao para a proposta."""

    numero: str = ""
    orgao: str = ""
    modalidade: str = ""
    objeto: str = ""
    data_abertura: str = ""


# ---------------------------------------------------------------------------
# Models — Planilha de Custos IN SEGES/ME 65/2021
# ---------------------------------------------------------------------------

class ItemCusto(BaseModel):
    """Item genérico de custo com descrição e valor."""

    descricao: str
    valor: Decimal = Decimal("0")
    referencia: str = ""  # ex: "CCT 2024", "SINAPI", "IN 65/2021"


class Modulo1Remuneracao(BaseModel):
    """Módulo 1 — Composição da Remuneração."""

    salario_base: Decimal = Decimal("0")
    adicional_periculosidade: Decimal = Decimal("0")
    adicional_insalubridade: Decimal = Decimal("0")
    adicional_noturno: Decimal = Decimal("0")
    hora_extra: Decimal = Decimal("0")
    outros: list[ItemCusto] = Field(default_factory=list)

    @property
    def total(self) -> Decimal:
        base = (
            self.salario_base + self.adicional_periculosidade
            + self.adicional_insalubridade + self.adicional_noturno
            + self.hora_extra
        )
        return base + sum(i.valor for i in self.outros)


class Modulo2Encargos(BaseModel):
    """Módulo 2 — Encargos e Benefícios Anuais, Mensais e Diários."""

    gps: Decimal = Decimal("0")  # 20%
    fgts: Decimal = Decimal("0")  # 8%
    salario_educacao: Decimal = Decimal("0")  # 2.5%
    sesi_sesc: Decimal = Decimal("0")  # 1.5%
    senai_senac: Decimal = Decimal("0")  # 1.0%
    sebrae: Decimal = Decimal("0")  # 0.6%
    incra: Decimal = Decimal("0")  # 0.2%
    rat: Decimal = Decimal("0")  # 1-3%
    vale_transporte: Decimal = Decimal("0")
    vale_alimentacao: Decimal = Decimal("0")
    assistencia_medica: Decimal = Decimal("0")
    seguro_vida: Decimal = Decimal("0")
    outros: list[ItemCusto] = Field(default_factory=list)

    @property
    def total(self) -> Decimal:
        encargos = (
            self.gps + self.fgts + self.salario_educacao
            + self.sesi_sesc + self.senai_senac + self.sebrae
            + self.incra + self.rat
        )
        beneficios = (
            self.vale_transporte + self.vale_alimentacao
            + self.assistencia_medica + self.seguro_vida
        )
        return encargos + beneficios + sum(i.valor for i in self.outros)


class Modulo3Rescisao(BaseModel):
    """Módulo 3 — Provisão para Rescisão."""

    aviso_previo_indenizado: Decimal = Decimal("0")
    incidencia_fgts_aviso: Decimal = Decimal("0")
    multa_fgts: Decimal = Decimal("0")
    aviso_previo_trabalhado: Decimal = Decimal("0")
    incidencia_fgts_trabalhado: Decimal = Decimal("0")
    outros: list[ItemCusto] = Field(default_factory=list)

    @property
    def total(self) -> Decimal:
        return (
            self.aviso_previo_indenizado + self.incidencia_fgts_aviso
            + self.multa_fgts + self.aviso_previo_trabalhado
            + self.incidencia_fgts_trabalhado
            + sum(i.valor for i in self.outros)
        )


class Modulo4Reposicao(BaseModel):
    """Módulo 4 — Custo de Reposição do Profissional Ausente."""

    ferias: Decimal = Decimal("0")
    ausencia_por_doenca: Decimal = Decimal("0")
    licenca_paternidade: Decimal = Decimal("0")
    ausencias_legais: Decimal = Decimal("0")
    ausencia_por_acidente: Decimal = Decimal("0")
    outros: list[ItemCusto] = Field(default_factory=list)

    @property
    def total(self) -> Decimal:
        return (
            self.ferias + self.ausencia_por_doenca
            + self.licenca_paternidade + self.ausencias_legais
            + self.ausencia_por_acidente
            + sum(i.valor for i in self.outros)
        )


class Modulo5Insumos(BaseModel):
    """Módulo 5 — Insumos Diversos."""

    uniformes: Decimal = Decimal("0")
    materiais: Decimal = Decimal("0")
    equipamentos: Decimal = Decimal("0")
    outros: list[ItemCusto] = Field(default_factory=list)

    @property
    def total(self) -> Decimal:
        return (
            self.uniformes + self.materiais + self.equipamentos
            + sum(i.valor for i in self.outros)
        )


class Modulo6CustosIndiretos(BaseModel):
    """Módulo 6 — Custos Indiretos, Lucro e Tributos."""

    custos_indiretos_pct: Decimal = Decimal("0")  # ex: 3%
    lucro_pct: Decimal = Decimal("0")  # ex: 6.79%
    pis: Decimal = Decimal("0.65")  # %
    cofins: Decimal = Decimal("3.0")  # %
    iss: Decimal = Decimal("5.0")  # %
    outros_tributos: list[ItemCusto] = Field(default_factory=list)

    @property
    def bdi_total_pct(self) -> Decimal:
        """BDI total como percentual."""
        return (
            self.custos_indiretos_pct + self.lucro_pct
            + self.pis + self.cofins + self.iss
            + sum(i.valor for i in self.outros_tributos)
        )

    def calcular_valor(self, subtotal: Decimal) -> Decimal:
        """Calcula valor do módulo 6 sobre o subtotal dos módulos 1-5."""
        return subtotal * self.bdi_total_pct / Decimal("100")


class PlanilhaCustos(BaseModel):
    """Planilha de Custos e Formação de Preços — IN SEGES/ME 65/2021.

    Para contratos de serviços continuados com dedicação exclusiva.
    """

    cargo: str
    quantidade_postos: int = 1
    modulo1: Modulo1Remuneracao = Field(default_factory=Modulo1Remuneracao)
    modulo2: Modulo2Encargos = Field(default_factory=Modulo2Encargos)
    modulo3: Modulo3Rescisao = Field(default_factory=Modulo3Rescisao)
    modulo4: Modulo4Reposicao = Field(default_factory=Modulo4Reposicao)
    modulo5: Modulo5Insumos = Field(default_factory=Modulo5Insumos)
    modulo6: Modulo6CustosIndiretos = Field(default_factory=Modulo6CustosIndiretos)

    @property
    def subtotal_modulos_1_5(self) -> Decimal:
        return (
            self.modulo1.total + self.modulo2.total
            + self.modulo3.total + self.modulo4.total
            + self.modulo5.total
        )

    @property
    def valor_modulo_6(self) -> Decimal:
        return self.modulo6.calcular_valor(self.subtotal_modulos_1_5)

    @property
    def valor_mensal_unitario(self) -> Decimal:
        return self.subtotal_modulos_1_5 + self.valor_modulo_6

    @property
    def valor_mensal_total(self) -> Decimal:
        return self.valor_mensal_unitario * self.quantidade_postos


# ---------------------------------------------------------------------------
# Generator
# ---------------------------------------------------------------------------

DECLARACAO_TIPOS_VALIDOS = ["me-epp", "fato-impeditivo", "menor", "independente"]


class PropostaGenerator:
    """Gera documentos DOCX a partir de templates."""

    def __init__(self, templates_dir: Path | None = None):
        self.templates_dir = templates_dir or TEMPLATES_DIR

    # ----- Proposta de Preços -----

    def gerar_proposta_preco(
        self,
        empresa: DadosEmpresa,
        licitacao: DadosLicitacao,
        itens: list[ItemProposta],
        output: Path,
        validade_dias: int = 60,
    ) -> Path:
        """Gera proposta de precos em DOCX."""
        template_path = self.templates_dir / "proposta_preco.docx"
        if not template_path.exists():
            return self._gerar_proposta_sem_template(empresa, licitacao, itens, output, validade_dias)

        doc = DocxTemplate(str(template_path))

        total_geral = sum(item.valor_total for item in itens)

        context = {
            "empresa": empresa.model_dump(),
            "licitacao": licitacao.model_dump(),
            "itens": [
                {
                    **item.model_dump(),
                    "valor_total": float(item.valor_total),
                    "valor_unitario": float(item.valor_unitario),
                    "quantidade": float(item.quantidade),
                }
                for item in itens
            ],
            "total_geral": float(total_geral),
            "validade_dias": validade_dias,
            "data_emissao": datetime.now().strftime("%d/%m/%Y"),
            "local_data": f"{datetime.now().strftime('%d de %B de %Y')}",
        }

        doc.render(context)
        doc.save(str(output))
        logger.info("Proposta gerada: %s", output)
        return output

    # ----- Declarações -----

    def gerar_declaracao(
        self,
        tipo: str,
        empresa: DadosEmpresa,
        output: Path,
    ) -> Path:
        """Gera declaracao padrao em DOCX.

        Tipos: me-epp, fato-impeditivo, menor, independente.
        """
        template_map = {
            "me-epp": "declaracao_me_epp.docx",
            "fato-impeditivo": "declaracao_fato_impeditivo.docx",
            "menor": "declaracao_menor.docx",
            "independente": "declaracao_independente.docx",
        }

        template_name = template_map.get(tipo)
        if not template_name:
            raise ValueError(f"Tipo de declaracao invalido: {tipo}. Validos: {list(template_map.keys())}")

        template_path = self.templates_dir / template_name
        if not template_path.exists():
            return self._gerar_declaracao_sem_template(tipo, empresa, output)

        doc = DocxTemplate(str(template_path))
        context = {
            "empresa": empresa.model_dump(),
            "data_emissao": datetime.now().strftime("%d/%m/%Y"),
            "local_data": f"{datetime.now().strftime('%d de %B de %Y')}",
        }
        doc.render(context)
        doc.save(str(output))
        logger.info("Declaracao gerada: %s", output)
        return output

    # ----- Planilha de Custos IN 65/2021 -----

    def gerar_planilha_custos(
        self,
        empresa: DadosEmpresa,
        licitacao: DadosLicitacao,
        planilhas: list[PlanilhaCustos],
        output: Path,
        vigencia_meses: int = 12,
    ) -> Path:
        """Gera planilha de custos e formação de preços (IN SEGES/ME 65/2021).

        Args:
            empresa: Dados da empresa.
            licitacao: Dados da licitação.
            planilhas: Uma PlanilhaCustos por cargo/posto.
            output: Caminho do arquivo de saída.
            vigencia_meses: Vigência do contrato em meses.
        """
        template_path = self.templates_dir / "planilha_custos.docx"
        if not template_path.exists():
            return self._gerar_planilha_sem_template(
                empresa, licitacao, planilhas, output, vigencia_meses
            )

        doc = DocxTemplate(str(template_path))
        context = self._contexto_planilha(empresa, licitacao, planilhas, vigencia_meses)
        doc.render(context)
        doc.save(str(output))
        logger.info("Planilha de custos gerada: %s", output)
        return output

    # ----- Checklist -----

    def gerar_checklist(
        self,
        licitacao: DadosLicitacao,
        edital: object | None = None,
    ) -> list[dict[str, str]]:
        """Gera checklist de documentação baseado nos dados do edital.

        Args:
            licitacao: Dados básicos da licitação.
            edital: EditalExtraido (opcional) para checklist mais preciso.

        Returns:
            Lista de dicts com 'documento', 'categoria' e 'obrigatorio'.
        """
        checklist: list[dict[str, str]] = []

        # Documentos sempre obrigatórios
        def _add(doc: str, cat: str) -> None:
            checklist.append({"documento": doc, "categoria": cat, "obrigatorio": "sim"})

        def _opt(doc: str, cat: str) -> None:
            checklist.append({"documento": doc, "categoria": cat, "obrigatorio": "condicional"})

        _add("Proposta de precos", "proposta")
        _add("Declaracao de habilitacao (Art. 63 Lei 14.133/21)", "declaracao")

        # Habilitação jurídica
        _add("Ato constitutivo / Contrato social", "juridica")
        _add("Documento de identidade do representante", "juridica")
        _add("Procuracao (se aplicavel)", "juridica")

        # Fiscal
        _add("Certidao conjunta de debitos federais (RFB/PGFN)", "fiscal")
        _add("Certidao negativa de debitos estaduais", "fiscal")
        _add("Certidao negativa de debitos municipais", "fiscal")
        _add("Certificado de regularidade FGTS (CRF)", "fiscal")
        _add("Certidao negativa de debitos trabalhistas (CNDT)", "fiscal")

        # Técnica
        _add("Atestado de capacidade tecnica", "tecnica")

        # Econômico-financeira
        _add("Certidao negativa de falencia/recuperacao judicial", "economico_financeira")

        # Declarações obrigatórias
        _add("Declaracao de nao emprego de menor", "declaracao")
        _add("Declaracao de inexistencia de fato impeditivo", "declaracao")
        _add("Declaracao de elaboracao independente de proposta", "declaracao")

        # Condicionais baseados no EditalExtraido
        if edital is not None:
            if getattr(edital, "exclusiva_me_epp", False):
                _add("Declaracao de ME/EPP (LC 123/2006)", "declaracao")

            if getattr(edital, "cota_reservada", False):
                _opt("Declaracao de ME/EPP para cota reservada", "declaracao")

            garantia = getattr(edital, "garantia", None)
            if garantia and getattr(garantia, "exige", False):
                pct = getattr(garantia, "percentual", "")
                _add(f"Garantia contratual ({pct}%)" if pct else "Garantia contratual", "garantia")

            if getattr(edital, "visita_tecnica", None) == "obrigatoria":
                _add("Atestado de visita tecnica", "tecnica")
            elif getattr(edital, "visita_tecnica", None) == "facultativa":
                _opt("Atestado de visita tecnica / Declaracao de dispensa", "tecnica")

            if getattr(edital, "exige_amostra", False):
                _add("Amostra do produto/material", "tecnica")

            # Habilitação detalhada
            hab = getattr(edital, "habilitacao_detalhada", None)
            if hab:
                if getattr(hab, "economico_financeira", []):
                    _add("Balanco patrimonial do ultimo exercicio", "economico_financeira")

        return checklist

    # ----- Fallbacks programáticos -----

    def _gerar_proposta_sem_template(
        self,
        empresa: DadosEmpresa,
        licitacao: DadosLicitacao,
        itens: list[ItemProposta],
        output: Path,
        validade_dias: int,
    ) -> Path:
        """Gera proposta programaticamente (sem template .docx)."""
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)

        doc.add_heading("PROPOSTA DE PRECOS", level=1)
        doc.add_paragraph(f"Razao Social: {empresa.razao_social}")
        doc.add_paragraph(f"CNPJ: {empresa.cnpj}")
        if empresa.endereco:
            doc.add_paragraph(f"Endereco: {empresa.endereco}")
        if empresa.email:
            doc.add_paragraph(f"E-mail: {empresa.email}")

        doc.add_heading("Referencia", level=2)
        doc.add_paragraph(f"Licitacao: {licitacao.numero}")
        doc.add_paragraph(f"Orgao: {licitacao.orgao}")
        doc.add_paragraph(f"Modalidade: {licitacao.modalidade}")
        doc.add_paragraph(f"Objeto: {licitacao.objeto}")

        doc.add_heading("Itens", level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["Item", "Descricao", "Un.", "Qtd.", "V. Unit.", "V. Total"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        total_geral = Decimal("0")
        for item in itens:
            row = table.add_row()
            row.cells[0].text = str(item.numero)
            row.cells[1].text = item.descricao
            row.cells[2].text = item.unidade
            row.cells[3].text = str(item.quantidade)
            row.cells[4].text = f"R$ {item.valor_unitario:,.2f}"
            row.cells[5].text = f"R$ {item.valor_total:,.2f}"
            total_geral += item.valor_total

        doc.add_paragraph(f"\nValor Total: R$ {total_geral:,.2f}")
        doc.add_paragraph(f"Validade: {validade_dias} dias")
        doc.add_paragraph(f"\nData: {datetime.now().strftime('%d/%m/%Y')}")

        if empresa.representante_legal:
            doc.add_paragraph(f"\n\n{empresa.representante_legal}")
            doc.add_paragraph(empresa.razao_social)

        doc.save(str(output))
        logger.info("Proposta gerada (sem template): %s", output)
        return output

    def _gerar_declaracao_sem_template(
        self,
        tipo: str,
        empresa: DadosEmpresa,
        output: Path,
    ) -> Path:
        """Gera declaracao programaticamente."""
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)

        textos = {
            "me-epp": (
                "DECLARACAO DE ENQUADRAMENTO COMO ME/EPP",
                f"{empresa.razao_social}, inscrita no CNPJ sob o no {empresa.cnpj}, "
                f"declara, sob as penas da lei, que se enquadra como Microempresa ou "
                f"Empresa de Pequeno Porte, nos termos da Lei Complementar no 123/2006, "
                f"e que nao se encontra em nenhuma das situacoes previstas no paragrafo "
                f"4o do artigo 3o da referida Lei.",
            ),
            "fato-impeditivo": (
                "DECLARACAO DE INEXISTENCIA DE FATO IMPEDITIVO",
                f"{empresa.razao_social}, inscrita no CNPJ sob o no {empresa.cnpj}, "
                f"declara, sob as penas da lei, que ate a presente data inexistem fatos "
                f"impeditivos para sua habilitacao no presente processo licitatorio, "
                f"ciente da obrigatoriedade de declarar ocorrencias posteriores.",
            ),
            "menor": (
                "DECLARACAO DE NAO EMPREGO DE MENOR",
                f"{empresa.razao_social}, inscrita no CNPJ sob o no {empresa.cnpj}, "
                f"declara, para fins do disposto no inciso XXXIII do art. 7o da "
                f"Constituicao Federal, que nao emprega menor de dezoito anos em "
                f"trabalho noturno, perigoso ou insalubre e nao emprega menor de "
                f"dezesseis anos, salvo na condicao de aprendiz, a partir de "
                f"quatorze anos.",
            ),
            "independente": (
                "DECLARACAO DE ELABORACAO INDEPENDENTE DE PROPOSTA",
                f"{empresa.razao_social}, inscrita no CNPJ sob o no {empresa.cnpj}, "
                f"declara, sob as penas da lei, que a proposta apresentada nesta "
                f"licitacao foi elaborada de maneira independente, e que o conteudo da "
                f"proposta nao foi, no todo ou em parte, direta ou indiretamente, "
                f"informado, discutido ou recebido de qualquer outro participante "
                f"potencial ou de fato desta licitacao, por qualquer meio ou pessoa.",
            ),
        }

        titulo, corpo = textos.get(tipo, ("DECLARACAO", "Texto padrao."))

        doc.add_heading(titulo, level=1)
        doc.add_paragraph(corpo)
        doc.add_paragraph(f"\nData: {datetime.now().strftime('%d/%m/%Y')}")

        if empresa.representante_legal:
            doc.add_paragraph(f"\n\n{empresa.representante_legal}")
        doc.add_paragraph(empresa.razao_social)
        doc.add_paragraph(f"CNPJ: {empresa.cnpj}")

        doc.save(str(output))
        logger.info("Declaracao %s gerada: %s", tipo, output)
        return output

    def _contexto_planilha(
        self,
        empresa: DadosEmpresa,
        licitacao: DadosLicitacao,
        planilhas: list[PlanilhaCustos],
        vigencia_meses: int,
    ) -> dict:
        """Monta contexto para template de planilha de custos."""
        cargos = []
        total_mensal = Decimal("0")
        for p in planilhas:
            cargos.append({
                "cargo": p.cargo,
                "postos": p.quantidade_postos,
                "mod1": float(p.modulo1.total),
                "mod2": float(p.modulo2.total),
                "mod3": float(p.modulo3.total),
                "mod4": float(p.modulo4.total),
                "mod5": float(p.modulo5.total),
                "mod6": float(p.valor_modulo_6),
                "mensal_unitario": float(p.valor_mensal_unitario),
                "mensal_total": float(p.valor_mensal_total),
            })
            total_mensal += p.valor_mensal_total

        total_global = total_mensal * vigencia_meses

        return {
            "empresa": empresa.model_dump(),
            "licitacao": licitacao.model_dump(),
            "cargos": cargos,
            "total_mensal": float(total_mensal),
            "total_global": float(total_global),
            "vigencia_meses": vigencia_meses,
            "data_emissao": datetime.now().strftime("%d/%m/%Y"),
        }

    def _gerar_planilha_sem_template(
        self,
        empresa: DadosEmpresa,
        licitacao: DadosLicitacao,
        planilhas: list[PlanilhaCustos],
        output: Path,
        vigencia_meses: int,
    ) -> Path:
        """Gera planilha de custos programaticamente."""
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(10)

        doc.add_heading("PLANILHA DE CUSTOS E FORMACAO DE PRECOS", level=1)
        doc.add_paragraph("IN SEGES/ME no 65/2021")
        doc.add_paragraph(f"Empresa: {empresa.razao_social} — CNPJ: {empresa.cnpj}")
        doc.add_paragraph(f"Licitacao: {licitacao.numero} — {licitacao.orgao}")
        doc.add_paragraph(f"Objeto: {licitacao.objeto}")
        doc.add_paragraph(f"Vigencia: {vigencia_meses} meses")
        doc.add_paragraph("")

        total_mensal_geral = Decimal("0")

        for p in planilhas:
            doc.add_heading(f"Cargo: {p.cargo} ({p.quantidade_postos} posto(s))", level=2)

            # Tabela de módulos
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.rows[0].cells[0].text = "Componente"
            table.rows[0].cells[1].text = "Valor (R$)"

            linhas = [
                ("Modulo 1 — Remuneracao", p.modulo1.total),
                ("  Salario base", p.modulo1.salario_base),
                ("Modulo 2 — Encargos e Beneficios", p.modulo2.total),
                ("Modulo 3 — Provisao Rescisao", p.modulo3.total),
                ("Modulo 4 — Reposicao Ausencia", p.modulo4.total),
                ("Modulo 5 — Insumos", p.modulo5.total),
                ("SUBTOTAL (Mod. 1-5)", p.subtotal_modulos_1_5),
                (f"Modulo 6 — Custos Indiretos/Lucro/Tributos ({p.modulo6.bdi_total_pct}%)", p.valor_modulo_6),
                ("VALOR MENSAL UNITARIO", p.valor_mensal_unitario),
                ("VALOR MENSAL TOTAL", p.valor_mensal_total),
            ]

            for desc, val in linhas:
                row = table.add_row()
                row.cells[0].text = desc
                row.cells[1].text = f"R$ {val:,.2f}"

            total_mensal_geral += p.valor_mensal_total
            doc.add_paragraph("")

        total_global = total_mensal_geral * vigencia_meses
        doc.add_heading("Resumo", level=2)
        doc.add_paragraph(f"Total Mensal: R$ {total_mensal_geral:,.2f}")
        doc.add_paragraph(f"Total Global ({vigencia_meses} meses): R$ {total_global:,.2f}")
        doc.add_paragraph(f"\nData: {datetime.now().strftime('%d/%m/%Y')}")

        if empresa.representante_legal:
            doc.add_paragraph(f"\n{empresa.representante_legal}")
        doc.add_paragraph(empresa.razao_social)

        doc.save(str(output))
        logger.info("Planilha de custos gerada: %s", output)
        return output
